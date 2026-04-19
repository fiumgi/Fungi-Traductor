import threading
import logging
from queue import Empty, Queue
from threading import Lock
from collections import defaultdict
import json
from pathlib import Path
from contextlib import contextmanager
from typing import Any, Literal

logger = logging.getLogger(__name__)


class TranslatorController:
    # Mapeo de códigos Argos (ISO 639-1) a Tesseract (ISO 639-2/específicos)
    _TESS_LANG_MAP = {
        "en": "eng", "es": "spa", "fr": "fra", "de": "deu",
        "it": "ita", "pt": "por", "ru": "rus", "zh": "chi_sim",
        "ja": "jpn", "ko": "kor", "ar": "ara", "hi": "hin",
        "tr": "tur", "nl": "nld", "pl": "pol", "uk": "ukr"
    }

    def __init__(self, view, model):
        self.view = view
        self.model = model
        self._pairs_by_source = {}
        self._suspend_auto = False
        self._ui_queue: Queue[tuple[Any, ...]] = Queue(maxsize=50)
        self._init_state = "idle"

        # Cache para evitar búsquedas repetidas
        self._voice_cache: list[tuple[str | None, str]] = []
        self._voice_cache_lang = None
        self._last_translate_text = ""
        self._translate_scheduled = False
        self._translate_timer = None
        self._current_translate_evt = None  # Evento para cancelar hilos antiguos

        # Lock para operaciones thread-safe
        self._pairs_lock = Lock()

        # Conectar botones
        self.view.btn_translate.config(command=self.translate)
        self.view.btn_detect.config(command=self._detect_language_async)
        self.view.btn_tts.config(command=self._text_to_speech_async)
        self.view.btn_swap_center.config(command=self.swap_languages)
        self.view.btn_clear.config(command=self.clear)
        self.view.btn_copy.config(command=self.copy)
        self.view.btn_open.config(command=self._on_open_file)
        self.view.btn_save.config(command=self._on_save_file)
        self.view.bind_auto_toggle(self.toggle_auto)
        self.view.bind_input_change(self._on_input_change)
        self.view.bind_from_change(self._on_from_change)
        self.view.bind_to_change(self._on_to_change)
        self.view.bind_voice_change(self._on_voice_change)
        self.view.bind_close(self.on_close)

    @contextmanager
    def _suspend_auto_translation(self):
        """Context manager para suspender la traducción automática durante cambios en UI."""
        self._suspend_auto = True
        try:
            yield
        finally:
            self._suspend_auto = False

    # ── INICIALIZACIÓN ─────────────────────────────────────────

    def _set_status(self, msg, level="info"):
        try:
            from queue import Full
            self._ui_queue.put_nowait(("status", msg, level))
        except Full:
            pass

    def _set_loading(
            self,
            active: bool,
            mode: Literal["determinate", "indeterminate"] = "indeterminate",
            value: float | None = None,
            detail: str = ""):
        try:
            from queue import Full
            self._ui_queue.put_nowait(("loading", active, mode, value, detail))
        except Full:
            pass

    def _apply_status(self, msg, level="info"):
        self.view.set_status(msg, level)

    def initialize(self):
        """Inicia el proceso de carga de paquetes e interfaz en segundo plano."""
        self._init_state = "loading"
        self._set_loading(
            True,
            mode="determinate",
            value=5,
            detail="Preparando inicio…")
        self._set_status("● actualizando índice de paquetes…", "info")
        self._load_config()
        self._schedule_ui_queue_poll()
        threading.Thread(target=self._initialize_async, daemon=True).start()

    def _initialize_async(self):
        # Inicializar modelo (Argos packages)
        if not self.model.init_packages(
                self._on_init_status,
                self._on_init_progress):
            try:
                from queue import Full
                self._ui_queue.put_nowait(("init_done", False))
            except Full:
                pass
            return

        # Poblar idiomas
        try:
            from queue import Full
            self._ui_queue.put_nowait(("populate",))
            self._ui_queue.put_nowait(("init_done", True))
        except Full:
            pass

    def _on_init_status(self, msg, level="info"):
        if level == "warn":
            self._init_state = "offline"
        elif level == "error":
            self._init_state = "error"
        self._set_status(msg, level)

    def _on_init_progress(self, _stage, value, detail):
        self._set_loading(True, mode="determinate", value=value, detail=detail)

    def _on_install_progress(self, _stage, value, detail):
        self._set_loading(True, mode="determinate", value=value, detail=detail)

    def _schedule_ui_queue_poll(self):
        # 50ms en lugar de 100ms para más fluidez
        self.view.after(50, self._drain_ui_queue)

    def _drain_ui_queue(self):
        try:
            # Procesar múltiples tareas en una sola pasada
            while True:
                task = self._ui_queue.get_nowait()
                kind = task[0]

                if kind == "status":
                    _, msg, level = task
                    self._apply_status(msg, level)
                elif kind == "loading":
                    _, active, mode, value, detail = task
                    self.view.set_loading(
                        active, mode=mode, value=value, detail=detail)
                elif kind == "populate":
                    self._populate_language_lists()
                elif kind == "init_done":
                    _, success = task
                    self._finish_initialize(success)
                elif kind == "toggle_ui":
                    _, enabled = task
                    self._toggle_ui(enabled)
                elif kind == "set_output":
                    _, result = task
                    self.view.set_output(result)
                elif kind == "select_from":
                    _, lang = task
                    self.view.select_from(lang)
                elif kind == "refresh_targets":
                    _, preferred = task
                    self._refresh_targets(preferred_code=preferred)
        except Empty:
            pass
        finally:
            self._schedule_ui_queue_poll()

    def _finish_initialize(self, success):
        self.view.set_loading(False)
        if not success:
            return

        if self._init_state == "offline":
            self._apply_status(
                "⚠ sin conexión — usando paquetes locales", "warn")
        elif self._init_state != "error":
            self._init_state = "ready"
            self._apply_status("● listo", "ok")

        self._toggle_ui(True)
        self._check_optional_deps()  # Volver a verificar tras habilitar UI

    def _check_optional_deps(self):
        """Verifica dependencias opcionales y deshabilita funciones si faltan."""
        # 1. Verificar TTS (pyttsx3)
        try:
            import pyttsx3
        except ImportError:
            self.view.set_button_enabled("btn_tts", False)
            self.view.set_tooltip("btn_tts", "Falta pyttsx3: pip install pyttsx3")
            logger.warning("TTS deshabilitado: pyttsx3 no instalado")

        # 2. Verificar Detección (langdetect)
        try:
            import langdetect
        except ImportError:
            self.view.set_button_enabled("btn_detect", False)
            self.view.set_tooltip("btn_detect", "Falta langdetect: pip install langdetect")
            logger.warning("Detección deshabilitada: langdetect no instalado")

        # 3. Verificar formatos de archivo (solo log y tooltips preventivos)
        missing_formats = []
        try:
            import fitz
        except ImportError:
            missing_formats.append("PDF")
        try:
            import docx
        except ImportError:
            missing_formats.append("Word (.docx)")
        try:
            import odf
        except ImportError:
            missing_formats.append("ODT")
        try:
            import pytesseract
            import PIL
        except ImportError:
            missing_formats.append("Imágenes (OCR)")

        if missing_formats:
            msg = f"Soporte limitado. Faltan: {', '.join(missing_formats)}"
            self.view.set_tooltip("btn_open", msg)
            logger.info(f"Soporte de archivos limitado: faltan {missing_formats}")

    def _get_config_path(self) -> Path:
        """Determina la ruta del archivo de configuración usando platformdirs."""
        try:
            from platformdirs import user_config_dir
            config_dir = Path(user_config_dir("FungiTraductor", "fiumgi"))
            config_dir.mkdir(parents=True, exist_ok=True)
            return config_dir / "config.json"
        except Exception:
            # Fallback al directorio del paquete si falla platformdirs
            return Path(__file__).parent.parent / "config.json"

    def _load_config(self):
        """Carga la configuración persistente (idiomas, auto-traducción, geometría)."""
        path = self._get_config_path()
        if not path.exists():
            return

        try:
            with open(path, "r", encoding="utf-8") as f:
                config = json.load(f)

            # Restaurar par de idiomas si están disponibles
            if "last_from" in config:
                self.view.select_from(config["last_from"])
                # Disparar actualización de destinos
                self._refresh_targets(preferred_code=config.get("last_to"))

            # Restaurar auto-traducción
            if config.get("auto_enabled"):
                self.view.set_auto(True)

            # Restaurar geometría de ventana
            if "geometry" in config:
                self.view.geometry(config["geometry"])

            logger.info(f"Configuración cargada desde {path}")
        except Exception as e:
            logger.error(f"Error al cargar configuración: {e}")

    def _save_config(self):
        """Guarda la configuración actual."""
        path = self._get_config_path()
        try:
            config = {
                "last_from": self.view.get_from_code(),
                "last_to": self.view.get_to_code(),
                "auto_enabled": self.view.auto_enabled,
                "geometry": self.view.geometry()
            }
            with open(path, "w", encoding="utf-8") as f:
                json.dump(config, f, indent=4)
            logger.info(f"Configuración guardada en {path}")
        except Exception as e:
            logger.error(f"Error al guardar configuración: {e}")

    def _populate_language_lists(self):
        """Versión optimizada con mejor gestión de memoria"""
        pairs = self.model.available_pairs()
        from_langs = {}
        pairs_by_source: dict[str, dict[str, str]] = defaultdict(dict)  # Más eficiente que setdefault

        for src_code, tgt_code, src_name, tgt_name in pairs:
            from_langs[src_code] = src_name
            pairs_by_source[src_code][tgt_code] = tgt_name

        # Cache directo sin lambda innecesarios
        new_pairs = {}
        for src_code, targets in pairs_by_source.items():
            new_pairs[src_code] = sorted(
                targets.items(),
                key=lambda item: (item[1].lower(), item[0])
            )

        with self._pairs_lock:
            self._pairs_by_source = new_pairs

        from_items = sorted(
            from_langs.items(),
            key=lambda item: (
                item[1].lower(),
                item[0]))
        
        self.view.populate_from(from_items)

        if not from_items:
            self.view.populate_to([])
            if self._init_state not in {"error", "offline"}:
                self._set_status("● no hay idiomas disponibles", "warn")
            return

        default_source = "es" if "es" in self._pairs_by_source else from_items[0][0]
        self.view.select_from(default_source)
        self._refresh_targets(preferred_code="en")

    def _refresh_targets(self, preferred_code=None):
        src = self.view.get_from_code()
        with self._pairs_lock:
            targets = self._pairs_by_source.get(src, [])

        self.view.populate_to(targets)
        if not targets:
            self._set_status(
                f"● no hay destinos disponibles para {src}", "warn")
            return

        current_code = self.view.get_to_code()
        available_codes = {code for code, _ in targets}

        if preferred_code in available_codes:
            selected_code = preferred_code
        elif current_code in available_codes:
            selected_code = current_code
        else:
            selected_code = targets[0][0]

        self.view.select_to(selected_code)
        self._refresh_voices()
        self._update_panel_headers()

    def _refresh_voices(self, preferred_voice_id=None):
        """Versión mejorada con caching de voces"""
        lang_code = self.view.get_to_code()

        # Cache de voces para evitar recargarlas si el idioma no cambió
        if self._voice_cache_lang != lang_code:
            self._voice_cache_lang = lang_code
            voice_items: list[tuple[str | None, str]] = [(None, "Automática")]
            voice_items.extend(self.model.list_voices(lang_code))
            self._voice_cache = voice_items

        self.view.populate_voices(self._voice_cache)

        current_voice_id = self.view.get_selected_voice_id()
        if preferred_voice_id is not None:
            self.view.select_voice(preferred_voice_id)
        elif current_voice_id is not None:
            self.view.select_voice(current_voice_id)
        else:
            self.view.select_voice(None)

    def _update_panel_headers(self):
        from_val = self.view.from_combo.get()
        to_val = self.view.to_combo.get()
        src_name = from_val.split(' (')[0] if from_val else "..."
        tgt_name = to_val.split(' (')[0] if to_val else "..."
        self.view.lbl_src.config(text=f"Texto original · {src_name}")
        self.view.lbl_dst.config(text=f"Traducción · {tgt_name}")

    def _on_from_change(self, _event=None):
        previous_target = self.view.get_to_code()
        self._refresh_targets(preferred_code=previous_target)
        if self.view.auto_enabled and self.view.get_input().strip():
            self._schedule_translate()

    def _on_to_change(self, _event=None):
        self._voice_cache_lang = None  # Invalidar cache de voces
        self._refresh_voices()
        self._update_panel_headers()
        if self.view.auto_enabled and self.view.get_input().strip():
            self._schedule_translate()

    def _on_voice_change(self, _event=None):
        if self.view.get_selected_voice_id():
            self._set_status("● voz manual seleccionada", "info")
        else:
            self._set_status("● voz automática activada", "info")

    def _on_input_change(self, _event=None):
        """Versión mejorada con debouncing para auto-traducción"""
        if self._suspend_auto:
            return

        text = self.view.get_input()
        self.view.set_char_count(len(text))

        if self.view.auto_enabled:
            if text.strip():
                # Solo traducir si el texto cambió y no hay traducción
                # pendiente
                if text != self._last_translate_text:
                    self._schedule_translate()
            else:
                self.view.set_output("")
                self._set_status("● limpio", "info")

    def _schedule_translate(self):
        """Debouncing: evita traducir en cada carácter"""
        if self._translate_timer is not None:
            self.view.after_cancel(self._translate_timer)

        self._translate_timer = self.view.after(300, self._execute_translate)

    def _execute_translate(self):
        """Ejecuta la traducción con debouncing"""
        self._translate_timer = None
        self.translate()

    def toggle_auto(self):
        """Alterna el modo de traducción automática."""
        enabled = not self.view.auto_enabled
        self.view.set_auto(enabled)
        self._set_status(
            f"● auto {
                'activado' if enabled else 'desactivado'}",
            "info")

        if enabled and self.view.get_input().strip():
            self._schedule_translate()

    # ── FUNCIONES ─────────────────────────────────────────────

    def translate(self):
        """Traducción en hilo separado para no bloquear UI"""
        # Cancelar cualquier temporizador automático pendiente
        if self._translate_timer is not None:
            self.view.after_cancel(self._translate_timer)
            self._translate_timer = None

        text = self.view.get_input()
        src = self.view.get_from_code()
        tgt = self.view.get_to_code()

        if not text.strip():
            return

        if len(text) > 5000:
            self._set_status(
                "● texto largo: la traducción puede tardar", "warn")

        self._toggle_ui(False)
        self._last_translate_text = text

        # Cancelar traducción en curso si existe
        if self._current_translate_evt:
            self._current_translate_evt.set()

        # Crear nuevo evento para esta traducción
        evt = threading.Event()
        self._current_translate_evt = evt

        # Ejecutar en hilo separado para no bloquear UI
        threading.Thread(
            target=self._translate_async,
            args=(text, src, tgt, evt),
            daemon=True
        ).start()

    def _translate_async(self, text, src, tgt, cancel_evt):
        """Traducción asincrónica en segundo plano"""
        try:
            if cancel_evt.is_set():
                return

            with self._pairs_lock:
                valid_targets = {
                    code for code,
                    _ in self._pairs_by_source.get(
                        src,
                        [])}
            if tgt not in valid_targets:
                self._set_status(
                    "● combinación de idiomas no disponible", "warn")
                return

            self._set_loading(
                True,
                mode="determinate",
                value=10,
                detail="Validando paquete…")
            try:
                if cancel_evt.is_set():
                    return
                if not self.model.ensure_pair(
                        src, tgt, self._set_status, self._on_install_progress):
                    return

                if cancel_evt.is_set():
                    return
                self._set_loading(
                    True,
                    mode="determinate",
                    value=90,
                    detail="Traduciendo texto…")
                result = self.model.translate(
                    text, src, tgt, on_progress=self._on_install_progress)

                if cancel_evt.is_set():
                    return
                self._ui_queue.put_nowait(("set_output", result))
                self._set_status("● traducción lista", "ok")
            finally:
                self._set_loading(False)
                try:
                    from queue import Full
                    self._ui_queue.put_nowait(("toggle_ui", True))
                except Full:
                    pass
        except Exception as e:
            self._set_loading(False)
            self._set_status(f"● error: {e}", "error")
            try:
                self._ui_queue.put_nowait(("toggle_ui", True))
            except Exception:
                pass

    def _toggle_ui(self, enabled: bool):
        """Habilita o deshabilita los controles principales de la interfaz"""
        state = "normal" if enabled else "disabled"
        self.view.btn_translate.config(state=state)
        self.view.btn_detect.config(state=state)
        self.view.btn_swap_center.config(state=state)
        self.view.btn_clear.config(state=state)
        self.view.btn_open.config(state=state)
        self.view.btn_tts.config(state=state)
        self.view.from_combo.config(
            state="readonly" if enabled else "disabled")
        self.view.to_combo.config(state="readonly" if enabled else "disabled")
        self.view.input_text.config(state=state)

    def _detect_language_async(self):
        """Detección en hilo separado"""
        text = self.view.get_input()
        if not text.strip():
            return

        current_target = self.view.get_to_code()
        self._toggle_ui(False)
        threading.Thread(
            target=self.detect_language,
            args=(text, current_target),
            daemon=True
        ).start()

    def detect_language(self, text: str, current_target: str):
        """Detecta el idioma del texto ingresado y actualiza el origen."""
        try:
            lang = self.model.detect(text)
            if not lang:
                self._set_status("● no se pudo detectar el idioma", "warn")
                return

            with self._pairs_lock:
                available = lang in self._pairs_by_source

            if not available:
                self._set_status(
                    f"● idioma detectado no disponible: {lang}", "warn")
                return

            try:
                self._ui_queue.put_nowait(("select_from", lang))
                self._ui_queue.put_nowait(("refresh_targets", current_target))
            except Exception:
                pass
            self._set_status(f"● detectado: {lang}", "info")
        except Exception as e:
            self._set_status(f"● error detección: {e}", "error")
        finally:
            self._ui_queue.put_nowait(("toggle_ui", True))

    def _text_to_speech_async(self):
        """TTS ya se ejecuta en hilo separado en el modelo, pero añadimos feedback"""
        text = self.view.get_output()
        if not text.strip():
            return

        self._set_status("● reproduciendo audio…", "info")
        threading.Thread(
            target=self.text_to_speech,
            daemon=True
        ).start()

    def text_to_speech(self):
        """Ejecuta la lectura por voz (TTS) del texto traducido."""
        text = self.view.get_output()
        lang = self.view.get_to_code()
        voice_id = self.view.get_selected_voice_id()

        if not text.strip():
            return

        try:
            self.model.speak(text, lang, voice_id)
        except Exception as e:
            self._set_status(f"● error TTS: {e}", "error")

    def swap_languages(self):
        """Intercambia el idioma de origen y destino, y sus textos correspondientes."""
        try:
            src = self.view.get_from_code()
            tgt = self.view.get_to_code()

            input_text = self.view.get_input()
            output_text = self.view.get_output()

            with self._suspend_auto_translation():
                self.view.select_from(tgt)
                self._voice_cache_lang = None  # Invalidar cache
                self._refresh_targets(preferred_code=src)

                self.view.set_input(output_text)
                self.view.set_output(input_text)
                self.view.set_char_count(len(output_text))

            if self.view.auto_enabled and output_text.strip():
                self._schedule_translate()

        except Exception as e:
            self._set_status(f"● error swap: {e}", "error")

    def clear(self):
        """Limpia los paneles de entrada y salida."""
        with self._suspend_auto_translation():
            self.view.set_input("")
            self.view.set_output("")
            self.view.set_char_count(0)
        self._set_status("● limpio", "info")

    def copy(self):
        """Copia el texto traducido al portapapeles."""
        text = self.view.get_output()
        if text.strip():
            self.view.clipboard_clear()
            self.view.clipboard_append(text)
            self.view.btn_copy.config(text="✓ Copiado", fg="#43e97b")
            self.view.after(
                1500,
                lambda: self.view.btn_copy.config(
                    text="Copiar", fg="#7a7f96"))
            self._set_status("● copiado al portapapeles", "ok")

    def _on_open_file(self):
        """Carga el contenido de un archivo (.txt, .pdf, .docx, .odt) en el input"""
        path = self.view.ask_open_file()
        if not path:
            return

        try:
            content = self._extract_text_from_file(path)
            if content is None:  # Error manejado internamente
                return

            with self._suspend_auto_translation():
                self.view.set_input(content)
                self.view.set_char_count(len(content))

            self._set_status(f"● archivo cargado: {len(content)} caracteres", "ok")

            if self.view.auto_enabled and content.strip():
                self._schedule_translate()
        except Exception as e:
            self._set_status(f"✗ error al leer archivo: {e}", "error")
            logger.error(f"Error al abrir archivo {path}: {e}")

    def _extract_text_from_file(self, path: str) -> str | None:
        """Extrae texto de un archivo según su extensión."""
        ext = path.lower()
        if ext.endswith(".pdf"):
            return self._extract_pdf(path)
        elif ext.endswith(".docx"):
            return self._extract_docx(path)
        elif ext.endswith(".odt"):
            return self._extract_odt(path)
        elif ext.endswith((".png", ".jpg", ".jpeg")):
            return self._extract_image_ocr(path)
        else:
            try:
                with open(path, "r", encoding="utf-8") as f:
                    return f.read()
            except UnicodeDecodeError:
                with open(path, "r", encoding="latin-1") as f:
                    return f.read()

    def _extract_pdf(self, path: str) -> str:
        """Extrae texto de un archivo PDF."""
        self._set_status("● extrayendo texto de PDF (preservando layout)…", "info")
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        content_parts = []
        for page in doc:
            blocks = page.get_text("blocks", sort=True)
            for b in blocks:
                text = b[4].strip()
                if text:
                    content_parts.append(text)
        doc.close()
        return "\n\n".join(content_parts)

    def _extract_docx(self, path: str) -> str:
        """Extrae texto de un archivo Word (.docx)."""
        self._set_status("● extrayendo texto de Word…", "info")
        import docx
        doc = docx.Document(path)
        return "\n".join([p.text for p in doc.paragraphs])

    def _extract_odt(self, path: str) -> str:
        """Extrae texto de un archivo OpenDocument (.odt)."""
        self._set_status("● extrayendo texto de ODT…", "info")
        from odf import opendocument, teletype
        doc = opendocument.load(path)
        return teletype.extractText(doc)

    def _extract_image_ocr(self, path: str) -> str | None:
        """Extrae texto de una imagen usando OCR."""
        self._set_status("● extrayendo texto de imagen (OCR)…", "info")
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(path)

            # Detectar idioma origen para mejorar precisión del OCR
            src_code = self.view.get_from_code()
            tess_lang = self._TESS_LANG_MAP.get(src_code, "eng")

            self._set_status(f"● OCR en progreso ({tess_lang})…", "info")
            content = pytesseract.image_to_string(img, lang=tess_lang).strip()
            
            if not content:
                self._set_status("✗ no se detectó texto en la imagen", "warn")
                return ""
            return content
        except ImportError:
            self._set_status("✗ faltan dependencias: pip install pytesseract Pillow", "error")
            return None
        except Exception as ex:
            logger.error(f"Error OCR: {ex}")
            self._set_status("✗ error OCR: ¿está instalado tesseract-ocr en el sistema?", "error")
            return None

    def _on_save_file(self):
        """Guarda la traducción actual en un archivo .txt"""
        content = self.view.get_output()
        if not content.strip():
            self._set_status("● nada que guardar", "warn")
            return

        path = self.view.ask_save_file()
        if not path:
            return

        try:
            ext = path.lower()
            if ext.endswith(".pdf"):
                from fpdf import FPDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Helvetica", size=12)
                # Intentar registrar fuente Unicode si está disponible en el sistema
                _unicode_font_paths = [
                    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                    "/usr/share/fonts/dejavu/DejaVuSans.ttf",
                    "C:/Windows/Fonts/arial.ttf",
                ]
                _unicode_font_loaded = False
                for _font_path in _unicode_font_paths:
                    import os as _os
                    if _os.path.exists(_font_path):
                        try:
                            pdf.add_font("Unicode", fname=_font_path)
                            pdf.set_font("Unicode", size=12)
                            _unicode_font_loaded = True
                        except Exception:
                            pass
                        break
                if not _unicode_font_loaded:
                    # Fallback: reemplazar caracteres no soportados con '?'
                    content = content.encode(
                        'windows-1252', 'replace').decode('windows-1252')
                pdf.multi_cell(0, 8, text=content)
                pdf.output(path)
            elif ext.endswith(".docx"):
                import docx
                doc = docx.Document()
                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
                doc.save(path)
            elif ext.endswith(".odt"):
                from odf.opendocument import OpenDocumentText
                from odf.text import P
                doc = OpenDocumentText()
                for line in content.split('\n'):
                    if line.strip():
                        doc.text.addElement(P(text=line))
                doc.save(path)
            else:
                # Fallback a texto plano
                if not ext.endswith(".txt"):
                    path += ".txt"
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)

            self._set_status(
                f"● traducción guardada como {ext.split('.')[-1].upper()}", "ok")
        except Exception as e:
            self._set_status(f"✗ error al guardar: {e}", "error")

    def on_close(self):
        """Maneja el cierre de la ventana"""
        # Guardar configuración antes de cerrar
        self._save_config()

        # Cancelar cualquier traducción en curso
        if self._current_translate_evt:
            self._current_translate_evt.set()

        # Detener cualquier timer de debouncing
        if self._translate_timer:
            self.view.after_cancel(self._translate_timer)

        self.view.destroy()
